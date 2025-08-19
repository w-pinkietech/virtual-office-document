#!/usr/bin/env python3
"""
Markdown契約書類をWord形式(.docx)に変換するスクリプト
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import markdown2

def clean_markdown_text(text):
    """Markdownの特殊文字を処理"""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    return text

def convert_markdown_to_docx(md_file_path, docx_file_path):
    """MarkdownファイルをWord文書に変換"""
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    doc = Document()
    
    font_name = 'メイリオ'
    
    lines = md_content.split('\n')
    current_list_level = 0
    
    for line in lines:
        line = line.rstrip()
        
        if not line:
            doc.add_paragraph()
            continue
        
        if line.startswith('# '):
            title = clean_markdown_text(line[2:])
            p = doc.add_heading(title, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(16)
                
        elif line.startswith('## '):
            title = clean_markdown_text(line[3:])
            p = doc.add_heading(title, level=2)
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(14)
                
        elif line.startswith('### '):
            title = clean_markdown_text(line[4:])
            p = doc.add_heading(title, level=3)
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(12)
                
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        elif line.startswith('- ') or line.startswith('* '):
            text = clean_markdown_text(line[2:])
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(text)
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(10.5)
                
        elif re.match(r'^\d+\.\s', line):
            text = clean_markdown_text(re.sub(r'^\d+\.\s', '', line))
            p = doc.add_paragraph(style='List Number')
            p.add_run(text)
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(10.5)
                
        elif line.startswith('□ ') or line.startswith('☑ '):
            checkbox = '□' if line.startswith('□') else '☑'
            text = clean_markdown_text(line[2:])
            p = doc.add_paragraph()
            p.add_run(f'{checkbox} {text}')
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(10.5)
            p.left_indent = Inches(0.5)
            
        elif re.match(r'^\(\d+\)', line):
            text = clean_markdown_text(line)
            p = doc.add_paragraph()
            p.add_run(text)
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(10.5)
            p.left_indent = Inches(0.5)
            
        elif line.startswith('|'):
            pass
            
        elif line.startswith('**') and line.endswith('**'):
            text = clean_markdown_text(line)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(10.5)
            
        elif '＿' in line and not line.startswith('#'):
            parts = line.split('＿')
            p = doc.add_paragraph()
            for i, part in enumerate(parts):
                if i > 0:
                    p.add_run('＿' * 10)
                if part:
                    run = p.add_run(clean_markdown_text(part))
                    run.font.name = font_name
                    run.font.size = Pt(10.5)
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(10.5)
                
        else:
            text = clean_markdown_text(line)
            if text:
                p = doc.add_paragraph()
                p.add_run(text)
                for run in p.runs:
                    run.font.name = font_name
                    run.font.size = Pt(10.5)
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    doc.save(docx_file_path)
    print(f"✅ 変換完了: {md_file_path} → {docx_file_path}")

def main():
    """メイン処理"""
    base_dir = Path('/home/kenta/pinkie/virtual-office')
    md_dir = base_dir / '10_legal-documents' / 'current'
    docx_dir = base_dir / '10_legal-documents' / 'word'
    
    docx_dir.mkdir(parents=True, exist_ok=True)
    
    md_files = [
        '会員規約.md',
        '郵便サービス契約書_法人.md',
        'バーチャルオフィス申込書.md'
    ]
    
    for md_file in md_files:
        md_path = md_dir / md_file
        docx_file = md_file.replace('.md', '.docx')
        docx_path = docx_dir / docx_file
        
        if md_path.exists():
            print(f"📄 変換中: {md_file}")
            convert_markdown_to_docx(str(md_path), str(docx_path))
        else:
            print(f"❌ ファイルが見つかりません: {md_path}")
    
    print("\n✨ すべての変換が完了しました！")
    print(f"📁 Word文書の保存先: {docx_dir}")

if __name__ == "__main__":
    main()